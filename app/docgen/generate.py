import openai
from openai import AzureOpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import requests
import os

# Azure OpenAI API details


def gpt4o_url(prompt):
    """
    Generate text using the Azure OpenAI GPT-40 model.
    """
    client = AzureOpenAI(
        api_key=api_key,
        api_version=api_version,
        base_url=f"{api_base}/openai/deployments/{model}"
    )

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Generate a professional and structured text document based on the user prompt."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000,
        temperature=0.7,
    )
    return response.choices[0].message.content

def generate_docx(content, file_name="generated_document.docx", font_type="Arial", font_size=12, font_color=(0, 0, 0), existing_doc=None):
    """
    Convert GPT-generated content into a .docx file with customizable font type, size, and color.
    If an existing_doc is provided, content is added to it.
    """
    if existing_doc:
        doc = Document(existing_doc)
    else:
        doc = Document()

    if not existing_doc:
        title = doc.add_heading('Generated Document', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Custom font settings
    def set_font(run, font_type, font_size, font_color):
        if run:  # Check if the run exists
            run.font.name = font_type
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*font_color)

    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.startswith("# "):
            heading = doc.add_heading(para.replace("# ", "").strip(), level=1)
            if heading.runs:  # Check if heading has runs
                set_font(heading.runs[0], font_type, font_size, font_color)
        elif para.startswith("## "):
            heading = doc.add_heading(para.replace("## ", "").strip(), level=2)
            if heading.runs:  # Check if heading has runs
                set_font(heading.runs[0], font_type, font_size, font_color)
        elif para.startswith("### "):
            heading = doc.add_heading(para.replace("### ", "").strip(), level=3)
            if heading.runs:  # Check if heading has runs
                set_font(heading.runs[0], font_type, font_size, font_color)
        elif para.startswith("|"):
            rows = [row.split('|')[1:-1] for row in para.split('\n') if row.startswith('|')]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Table Grid')
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = cell.strip()
        elif para.startswith("!["):
            try:
                alt_text, image_path = para[2:].split("](")
                image_path = image_path.strip(")")
                add_image_to_doc(doc, image_path)
            except Exception as e:
                print(f"Error processing image line: {para} - {e}")
        else:
            paragraph = doc.add_paragraph(para.strip())
            if paragraph.runs:  # Check if the paragraph has runs
                set_font(paragraph.runs[0], font_type, font_size, font_color)

    doc.save(file_name)
    print(f"Document successfully saved as: {file_name}")


def add_image_to_doc(doc, image_path):
    """
    Adds an image to the document. Supports local paths and URLs.
    """
    try:
        if image_path.startswith("http"):
            response = requests.get(image_path)
            if response.status_code == 200:
                temp_image_path = "temp_image.jpg"
                with open(temp_image_path, "wb") as img_file:
                    img_file.write(response.content)
                doc.add_picture(temp_image_path, width=Inches(4))
                os.remove(temp_image_path)
            else:
                print(f"Failed to download image from URL: {image_path}")
        else:
            doc.add_picture(image_path, width=Inches(4))
    except Exception as e:
        print(f"Error adding image: {e}")

def convert_docx_to_pdf(docx_file, pdf_file, font_type="Arial", font_size=12, font_color=(0, 0, 0)):
    """
    Convert a .docx file to a .pdf using FPDF, with specified font type, size, and color.
    """
    doc = Document(docx_file)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font(font_type, size=font_size)

    # Convert font color to FPDF-compatible format
    font_color_r, font_color_g, font_color_b = font_color
    pdf.set_text_color(font_color_r, font_color_g, font_color_b)

    for paragraph in doc.paragraphs:
        if len(paragraph.runs) > 0:  # Check if the paragraph contains runs
            # Iterate through runs to get the text
            text = ""
            for run in paragraph.runs:
                text += run.text
            pdf.multi_cell(0, 10, text)
        else:
            # If no runs, just use the paragraph text
            pdf.multi_cell(0, 10, paragraph.text)

    pdf.output(pdf_file)
    print(f"Document successfully converted to PDF: {pdf_file}")

if __name__ == "__main__":
    prompt = """
    Create a simple document with:
   
    - lessons from life of Joseph.
    - add relevant url images
    - An image placeholder using this URL: https://lighteducation.pythonanywhere.com/tailwindUI/assets/images/my-avatar2.png
    - Proper headings and a paragraph of text
    """

    print("Generating content from GPT-40...")
    generated_content = gpt4o_url(prompt)
    print("\nGenerated Content:\n", generated_content)

    docx_file_name = "azure_openai_generated.docx"
    pdf_file_name = docx_file_name.replace(".docx", ".pdf")

    existing_document = None

    generate_docx(
        generated_content,
        file_name=docx_file_name,
        font_type="Times New Roman",
        font_size=14,
        font_color=(123, 45, 0),
        existing_doc=existing_document
    )

    convert_docx_to_pdf(
        docx_file_name,
        pdf_file_name,
        font_type="Arial",
        font_size=12
    )

    print("\nProcess completed. Check your working directory for the files.")
